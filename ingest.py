import os
import glob
from dotenv import load_dotenv
from agno.knowledge.pdf import PDFKnowledgeBase, PDFReader
from agno.knowledge.document import DocumentKnowledgeBase, Document
from agno.knowledge.combined import CombinedKnowledgeBase
from agno.vectordb.chroma import ChromaDb
from agno.embedder.huggingface import HuggingfaceCustomEmbedder
from docx import Document as DocxDocument

load_dotenv()

DATA_DIR    = "data"
CHROMA_PATH = "chroma_db"
COLLECTION  = "support_docs"

# ✅ HuggingFace embedder – using a lightweight public model
embedder = HuggingfaceCustomEmbedder("sentence-transformers/all-MiniLM-L6-v2")

# ✅ Vector DB setup
vector_db = ChromaDb(
    path=CHROMA_PATH,
    collection=COLLECTION,
    persistent_client=True,
    embedder=embedder
)

# ✅ Load and chunk PDFs (metadata will be auto-extracted from file names by the loader)
pdf_kb = PDFKnowledgeBase(
    path=DATA_DIR,
    reader=PDFReader(chunk=True),  # chunks and adds metadata automatically
    vector_db=vector_db,
    embedder=embedder,
)

# ✅ Parse DOCX files manually with required metadata
def load_docx_documents(folder: str):
    docs = []
    for path in glob.glob(os.path.join(folder, "**/*.docx"), recursive=True):
        try:
            docx_obj = DocxDocument(path)
            full_text = "\n".join([p.text for p in docx_obj.paragraphs])
            docs.append(Document(
                content=full_text,
                name=os.path.basename(path),
                meta_data={"source": os.path.basename(path)}  # ✅ important
            ))
        except Exception as e:
            print(f"⚠️ Skipped {path}: {e}")
    return docs

doc_kb = DocumentKnowledgeBase(
    documents=load_docx_documents(DATA_DIR),
    vector_db=vector_db,
    embedder=embedder,
)

# ✅ Combine KBs
knowledge_base = CombinedKnowledgeBase(
    sources=[pdf_kb, doc_kb],
    vector_db=vector_db,
    embedder=embedder,
)

# ✅ Ingest
if __name__ == "__main__":
    if not glob.glob(os.path.join(DATA_DIR, "*")):
        raise SystemExit("❌ No PDFs or DOCX found in ./data")

    knowledge_base.load(recreate=True)
    print("✅ Ingestion complete using HuggingFace embeddings.")
